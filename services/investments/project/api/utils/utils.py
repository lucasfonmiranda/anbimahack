from docx import Document
from docx.shared import Inches, Pt

def create_document(data_cessionario, data_cedente, data_client):
    cessionario = data_cessionario
    cedente = data_cedente
    client = data_client

    document = Document()
    
    header = document.add_heading('Anexo I – Carta do cliente', 0)
    header.alignment = 1

    date = document.add_paragraph('São Paulo, 01 de Setembro de 2019.')
    date.alignment = 2
    h = document.add_heading('À,', level=2)
    

    document.add_paragraph(cessionario['data']['name'])
    document.add_paragraph(cessionario['data']['address'])
    document.add_paragraph(cessionario['data']['cnpj'])
    document.add_paragraph(cessionario['data']['email'])

    document.add_paragraph("Com cópia para:")

    document.add_paragraph(cedente['data']['cedente'])
    document.add_paragraph(cedente['data']['cnpj'])
    document.add_paragraph(cedente['data']['email'])

    denominacao_distribuidor_cedente = cedente['data']['cedente']
    cnpj_distribuidor_cedente = cedente['data']['cnpj']



    ref = "Ref.: Transferência das aplicações em fundos de investimento na modalidade conta e ordem do {} CNPJ: {} ('Distribuidor Cedente')".format(denominacao_distribuidor_cedente, cnpj_distribuidor_cedente)
    p = document.add_paragraph()
    p.add_run(ref).bold = True


    document.add_heading('Prezados senhores,')
    p_four = document.add_paragraph('Venho, pela presente, solicitar e autorizar que o distribuidor cessionário receba as minhas aplicações em fundo de investimentos por meio do distribuidor cedente. Segue o discriminativo da minha posição em fundos de investimento que deverá ser recebida pelo distribuidor cessionário.').paragraph_format
    p_four.alignment = 3
    p_four.space_before
    records = (
        (3, '101', 'Spam'),
        (7, '422', 'Eggs'),
        (4, '631', 'Spam, spam, eggs, and spam')
    )

    table = document.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'NOME DO FUNDO'
    hdr_cells[1].text = 'CNPJ DO FUNDO '
    hdr_cells[2].text = 'QUANTIDADE DE COTAS'
    for qty, id, desc in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(qty)
        row_cells[1].text = id
        row_cells[2].text = desc
        
    document.add_paragraph()
    p_two = document.add_paragraph("Tenho ciência de que o distribuidor cessionário não tem nenhuma responsabilidade por atos ou fatos, de qualquer natureza, inclusive os relacionados às retenções e ao recolhimento do Imposto de Renda, e os obrigacionais, notadamente os relacionados no art. 33, da Instrução Normativa CVM 555/14, decorrentes das atividades relacionadas à distribuição dos fundos realizada pelo distribuidor cedente, e que, por conseguinte, qualquer reinvindicação relacionada ao período anterior à efetiva transferência dos meus investimentos para o distribuidor cessionário deverá ser direcionada ao distribuidor cedente.").paragraph_format
    p_two.alignment = 3
    p_two.space_before
    p_three = document.add_paragraph("Tenho ciência, ainda, de que os investimentos nos fundos, após a transferência, continuarão sendo realizados no formato “por conta e ordem” por meio do distribuidor cessionário, em consonância com o estabelecido na Instrução Normativa CVM 555, de 17 de dezembro de 2014.").paragraph_format
    p_three.alignment = 3
    p_three.space_before

    sigin = document.add_paragraph("_______________________________________________")
    sigin.alignment = 1

    client_name = client['data']['name']
    client_para = document.add_paragraph(client_name)
    client_para.alignment = 1

    cpf_client = client['data']['cpf']
    phase = "CPF: {}".format(cpf_client)
    cpf_para = document.add_paragraph(phase)
    cpf_para.alignment = 1


    document.save('documento.docx')
    return
